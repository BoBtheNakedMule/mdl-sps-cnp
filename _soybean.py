#pyinstaller C:\SPS_python\.venv\Marshall\_nifa-afri.py --onefile
import openpyxl
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT


from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog
import os

#TODO
#Add error handling for row misalignment
#import sys?


# Open File Prompt
def open_file():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(
        filetypes=[("Excel files", "*.xlsx")],
        defaultextension=".docx",
        title="SPS Current and Pending -NIFA-AFRI",
        initialdir=r"K:\_DeptAll\PreAward\3. Administrative\Faculty Documents - CPs\Faculty Current & Pending"
    )

    if not file_path:
        print("No File Selected")
        os.system="pause"
        exit()
    else:
        workbook_file_path = os.path.basename(file_path)

        return openpyxl.load_workbook(file_path), workbook_file_path

# Save File Prompt
def save_file(workbook_file_path):
    workbook_file_path = os.path.splitext(workbook_file_path)[0]
    save_path = filedialog.asksaveasfilename(
        filetypes=[("Word files", "*.docx")],
        title="SPS Current and Pending -NIFA-AFRI",
        initialdir=r"K:\_DeptAll\PreAward\3. Administrative\Faculty Documents - CPs\Faculty Current & Pending",
        initialfile=f"{workbook_file_path}.docx"
    )
    
    if not save_path:
        print("No File Saved")
        os.system('pause')
        exit()

    # Check if the file has a .docx extension, if not, add it
    elif not save_path.lower().endswith('.docx'):
        save_path += '.docx'
        save_file_error_handling(save_path)
    else:
        save_file_error_handling(save_path)
        

def save_file_error_handling(save_path):
    try:
        doc.save(save_path)
        print(f"Data has been exported to {save_path}")
        os.system('pause')

    except PermissionError:
        print("ERROR: File is open, close the file and try again")
        os.system('pause')

#function to handle the unusual borders on the main table
def set_cell_border(cell, **kwargs):

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for side, attrs in kwargs.items():
        pBdr = OxmlElement('w:tcBorders')
        tcPr.append(pBdr)
        
        border = OxmlElement(f'w:{side}')
        pBdr.append(border)

        for key, value in attrs.items():
            border.set(qn(f'w:{key}'), str(value))


# Main code
version = "SOYBEAN-CORN-20250801"
print(20 * "*")
print(f"Version: {version}")
print("When reporting issues, please provide this version number")
print(20 * "*")

# Load the Excel file
workbook, workbook_file_path = open_file()
sheet = workbook['C&P']

# Get the headers from row 40
headers = [cell.value for cell in sheet[40]]

# Create lists for active and pending projects
active_projects = []
pending_projects = []

# Iterate through rows starting from row 41
for row in sheet.iter_rows(min_row=41, values_only=True):
    nifa_category = row[headers.index('NIFA/AFRI/USDA Category')]
    if nifa_category == 'Active':
        active_projects.append((
            row[headers.index('NIFA/AFRI/USDA Person Name')],
            row[headers.index('NIFA/AFRI/USDA Agency Source')],
            row[headers.index('Total Award Amount (including Indirect Costs):  ')],
            row[headers.index('Project Period')],
            row[headers.index('Percentage of Time Committed')],
            row[headers.index('Project Title')]
        ))
    elif nifa_category == 'Pending':
        pending_projects.append((
            row[headers.index('NIFA/AFRI/USDA Person Name')],
            row[headers.index('NIFA/AFRI/USDA Agency Source')],
            row[headers.index('Total Award Amount (including Indirect Costs):  ')],
            row[headers.index('Project Period')],
            row[headers.index('Percentage of Time Committed')],
            row[headers.index('Project Title')]
        ))

# Function to create table in Word document
def create_table(doc, projects, title, category):

    #formats funding amount with dollar sign and two decimals. Handles error if the funding column has text
    def currency_formatting(funding_column):
        try:
            funding_float = float(funding_column if funding_column is not None else '')
            return f'${funding_float:,.0f}'
        except ValueError:
            print(f"The Funding Number is NOT text and says {funding_column}.\n Edit the Word file after saving or fix the Excel file and run this script again.")
            return '**** Incorrect Entry- Must be in format of ####.## ****'
    
    #Error checking to ensure other columns don't have blank spaces in needed rows.    
    def other_column_blank(data, label):
        if data is not None:
            return data
        else:
            print(f"Something is blank that shouldn't be, check {label} column")
            print("No File Saved")
            os.system(command="pause")
            exit() 
            
    table = doc.add_table(rows=2, cols=6)
    table.autofit = True
    table.style = 'Plain Table 4'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set custom borders for all cells
     
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # Top border
            if i == 0:
                set_cell_border(
                    cell, top={"sz": 8, "val": "double", "color": "000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell, top={"sz": 4, "val": "single", "color": "000000", "space": "0"}
                )
            # Bottom border
            if i == n_rows - 1:
                set_cell_border(
                    cell, bottom={"sz": 8, "val": "double", "color": "000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell, bottom={"sz": 4, "val": "single", "color": "000000", "space": "0"}
                )
            # Left border
            if j == 0:
                set_cell_border(
                    cell, left={"sz": 8, "val": "double", "color": "000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell, left={"sz": 4, "val": "single", "color": "000000", "space": "0"}
                )
            # Right border
            if j == n_cols - 1:
                set_cell_border(
                    cell, right={"sz": 8, "val": "double", "color": "000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell, right={"sz": 4, "val": "single", "color": "000000", "space": "0"}
                )
 
    
    # Set header row
    hdr_cells = table.rows[0].cells
    headers = ['NAME(S) OF PRINCIPAL INVESTIGATORS', 'SUPPORTING AGENCY', 'TOTAL $ AMOUNT', 'EFFECTIVE AND EXPIRATION DATES', ' % OF TIME COMMITTED', 'TITLE OF PROJECT']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set category row
    category_cells = table.rows[1].cells
    category_cells[1].text = f"{category}\n"
    for i in range(2, 6):
        category_cells[i].text = ""
    
    # Make the category row bold
    for cell in table.rows[1].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    
    # Add project data
    for index, (person, agency, amount, dates, time, title) in enumerate(projects):
        row_cells = table.add_row().cells
        for j, cell in enumerate(row_cells):
            # Set custom borders for new rows
            set_cell_border(
                cell,
                left={"sz": 0, "val": "double", "color": "#000000", "space": "0"},
                top={"sz": 0, "val": "none", "color": "#000000", "space": "0"}
            )
            # Add right border for all cells except the last column
            if j < len(row_cells) - 1:
                set_cell_border(
                    cell,
                    right={"sz": 4, "val": "single", "color": "#000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell,
                    right={"sz": 0, "val": "double", "color": "#000000", "space": "0"}
                )
            
            # Add bottom border only for the last row
            if index == len(projects) - 1:
                set_cell_border(
                    cell,
                    bottom={"sz": 4, "val": "double", "color": "#000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell,
                    bottom={"sz": 0, "val": "single", "color": "#000000", "space": "0"}
                )
        #Add Data to single row
        row_cells[0].text = str(other_column_blank(person, "PI Name"))
        row_cells[1].text = str(other_column_blank(agency, "NIFA Agency Source"))
        row_cells[2].text = str(currency_formatting(amount))
        row_cells[3].text = str(other_column_blank(dates, "Project Period"))
        row_cells[4].text = str(other_column_blank(time, "Percentage of Time Committed"))
        row_cells[5].text = str(other_column_blank(title, "Title"))




    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #cell.width = Inches(1.12)
                

'''def afri_parapgraph_check():
    

    while True:
        sponsor_type= input("Is this for AFRI (A) or NIFA (N) -Please answer with just one letter: ").lower()
        if sponsor_type == 'a' :
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("As an addendum to the Current and Pending Support, provide a brief summary below for any completed, current, or pending projects that appear similar to the current application, especially previous National Research Initiative (NRI) or AFRI awards: None")
            run.font.name = "Calibri (Body)"
            run.font.size = Pt(11)
            break #exit loop
             
        elif sponsor_type == 'n':
            break #exit loop

        else:
            print("*" * 20)
            print("INVALID ENTRY")
            sponsor_type = input("Please enter A for Afri or N for NIFA ").lower()'''
            


def create_document():
    # Create a new Word document
    document = Document()
    section = document.sections[0]
    section.orientation = 1
    section.page_width, section.page_height = section.page_height, section.page_width

    def add_formatted_paragraph(cell, text, font_size=12, bold=False, bullet=False, indent=0):
        paragraph = cell.add_paragraph(text, "List Number" if bullet else None)
        paragraph.paragraph_format.left_indent = Inches(indent)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in paragraph.runs:
            run.font.name = 'Calibri (Body)'
            run.font.size = Pt(font_size)
            run.bold = bold
    
    styles = document.styles
    style = styles.add_style('Plain Table 4', WD_STYLE_TYPE.TABLE)
    # style = document.styles['Normal'] to remove
    style.font.name = 'Calibri (Body)'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add a title to the document
 
    topline = document.add_paragraph()
    run = topline.add_run("CURRENT AND PENDING RESEARCH SUPPORT DISCLOUSRE")
    run.bold = True
    second_text = document.add_paragraph()
    run = second_text.add_run("FY 2026 Research Grant Application")  
    topline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a table with 1 row and 1 column
    header_table = document.add_table(rows=1, cols=1) 
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER   
    header_table.autofit = True
    '''for column in header_table.columns:
        for cell in column.cells:
            cell.width = Inches(6.7) '''

    # Get the cell and add the text
    cell = header_table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    #run = paragraph.add_run("\nCURRENT AND PENDING RESEARCH SUPPORT DISCLOUSRE")
    #run = paragraph.add_run("\nFY 2026 Reserch Grant Application")

    set_cell_border(
                    cell,
                    top={"sz": 18, "val": "none", "color": "#000000", "space": "0"},
                    bottom={"sz": 18, "val": "none", "color": "#000000", "space": "0"},
                    left={"sz": 18, "val": "none", "color": "#000000", "space": "0"},
                    right={"sz": 18, "val": "none", "color": "#000000", "space": "0"}
                )

    # Center the text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set the font to Calibri (Body), 12pt
    font = run.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)
    font.bold = False

    info_table = document.add_table(rows=1, cols=1)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Name row
    '''pi_name = sheet.cell(4,2).value.split(": ")[1]
    add_formatted_paragraph(info_table.cell(0, 0), f"Name: {pi_name}", font_size=12, bold=True)'''

    # Instructions row
    sec_cell = info_table.cell(0, 0)
   
    add_formatted_paragraph(sec_cell, "Instructions: ", bold=True)

    bullet_points = [
        "Record information for all active and pending projects/proposals, as well as all proposed research that is being considered by, or will be submitted to, 	other possible sponsors soon. (Concurrent submission of a proposal to other organizations will not prejudice its review.)",
        "All current research to which principal investigator(s) and other senior personnel have committed a portion of their time must be listed, whether salary/benefits for the person involved are included in the budgets of those projects.",
        "Use additional sheets if necessary."
    ]

    for point in bullet_points:
        add_formatted_paragraph(sec_cell, point, bullet=True, indent=0.5)


    # Get the second row of the info_table so you can set the border correctly
    #second_row = info_table.rows[0]

    # Add top border to each cell in the second row

    '''for cell in second_row.cells:
        set_cell_border(
            cell,
            top={"sz": 4, "val": "none", "color": "#FFFFFF", "space": "0"},
            left={"sz": 4, "val": "double", "color": "#FFFFFF", "space": "0"},
            right={"sz": 4, "val": "none", "color": "#FFFFFF", "space": "0"},
            bottom={"sz": 4, "val": "none", "color": "#FFFFFF", "space": "0"}
        )'''

    return document

# Create Word document and add tables
doc = create_document()

#TODO add if statement to handle empty projects
create_table(doc, active_projects, 'Active Projects', 'Current')

create_table(doc, pending_projects, 'Pending Projects', 'Pending')

#Add disclaimer if this is for AFRI
#afri_parapgraph_check()

print("*" * 20)
print("~~~After the Word file is created, be sure to adjust the column widths~~~")
print("*" * 20)
# Save the Word document

save_file(workbook_file_path)
