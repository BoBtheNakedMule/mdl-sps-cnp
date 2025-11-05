#pyinstaller C:\SPS_python\.venv\Marshall\_nasa_v3.py --onefile

import shutil
import docx
import os
from sys import exit
import shutil
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog


# Open File Prompt
def open_file():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    root.attributes('-topmost', True) #brings save dialog to the top
    file_path = filedialog.askopenfilename(
        filetypes=[("Excel files", "*.xlsx")],
        defaultextension=".xlsx",
        title="SPS Current and Pending -NASA",
        initialdir=r"K:\_DeptAll\PreAward\3. Administrative\Faculty Documents - CPs\Faculty Current & Pending"
    )

    if not file_path:
        print("No File Selected")
        os.system="pause"
        exit()
    else:
        workbook_file_path = os.path.basename(file_path)
        return openpyxl.load_workbook(file_path), workbook_file_path

# Save File Prompt
def save_file(workbook_file_path):
    workbook_file_path = os.path.splitext(workbook_file_path)[0]
    save_path = filedialog.asksaveasfilename(
        filetypes=[("Word files", "*.docx")],
        title="SPS Current and Pending NASA",
        initialdir=r"K:\_DeptAll\PreAward\3. Administrative\Faculty Documents - CPs\Faculty Current & Pending",
        initialfile=f"{workbook_file_path}.docx"
    )
    
    if not save_path:
        print("No File Saved")
        os.system('pause')
        exit()

    # Check if the file has a .docx extension, if not, add it
    elif not save_path.lower().endswith('.docx'):
        save_path += '.docx'
        save_file_error_handling(save_path)
    else:
        save_file_error_handling(save_path)
        

def save_file_error_handling(save_path):
    try:
        doc.save(save_path)
        print(f"Data has been exported to {save_path}")
        os.system('pause')

    except PermissionError:
        print("ERROR: File is open, close the file and try again")
        os.system('pause')


# Create a dictionary of column names and their corresponding column index
def create_column_dict(worksheet, row):
    column_dict = {}
    index = 0
    for cell in worksheet[row]:
        column_dict[cell.value] = index
        index+=1
    return column_dict


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    
    :param paragraph: The paragraph to which the hyperlink will be added.
    :param url: The URL for the hyperlink.
    :param text: The visible text for the hyperlink.
    """

    # Create a relationship ID for the hyperlink
    part = paragraph._parent.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run for the text
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set font to Times New Roman
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Times New Roman")
    font.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(font)

    # Set font size to 12 pt
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # 24 half-points = 12 points
    rPr.append(sz)

    # Apply the "Hyperlink" style (Word's default hyperlink appearance)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    run.append(rPr)

    # Create the text element and add it to the run
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)

    # Append the run to the hyperlink element
    hyperlink.append(run)

    # Append the hyperlink element to the paragraph
    paragraph._element.append(hyperlink)

# Function to create a table
def create_table(document, rows, table_title):
    
    # Function to set font for a paragraph
    def set_font(paragraph, is_left_column=False):
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if is_left_column:
                run.bold = True

    def add_table_row(table, label, value, is_bold=False):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value) if value is not None else ''
        set_font(cells[0].paragraphs[0], is_left_column=True)
        set_font(cells[1].paragraphs[0])
        if is_bold:
            cells[1].paragraphs[0].runs[0].bold = True

    #formats funding amount with dollar sign and two decimals. Handles error if the funding column has text
    def currency_formatting(funding_column,row):
        try:
            funding_float = float(funding_column if funding_column is not None else '')
            return f'${funding_float:,.2f}'
        except ValueError:
            print(f"The Funding Number is NOT text and says {funding_column} for title: {row}.\nEdit the Word file after saving or fix the Excel file and run this script again.")
            return '**** Incorrect Entry- Must be in format of ####.## ****'


    def blank_other_check(column, label):
        if column is None:
            print(f"Something is blank that shouldn't be, check {label} column")
            print("No File Saved")
            os.system(command="pause")
            exit()
        elif column is not None and label == "Overlap":
            overlap_column = f"{column}\n\n"
            return overlap_column
        else:
            return column

    title_formatting = document.add_paragraph()
    run = title_formatting.add_run(table_title)
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.bold = True


    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    #table.style = 'Plain Table 4'

    #handles inserting tables into NASA Doc only
    '''if insert_paragraph:
        insert_paragraph._element.addnext(table._element)
    else:
        document.add_paragraph()  # Add an empty paragraph for spacing
        document.add_table(table)'''

    #Set the width of the left column to 25%
    table.columns[0].width = Inches(1.5)
    table.columns[0].cells[0]._element.tcPr.tcW.type = 'pct'
    table.columns[0].cells[0]._element.tcPr.tcW.w = 1250  # 25% of 5000
    table.columns[1].width = Inches(8)
    # Set the preferred width of the right column to 75%
    table.columns[1].width = Inches(4.5)
    table.columns[1].cells[0]._element.tcPr.tcW.type = 'pct'
    table.columns[1].cells[0]._element.tcPr.tcW.w = 3750  # 75% of 5000

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = document.styles['Normal']
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for row in rows:
        add_table_row(table, 'Title:', row[1])
        add_table_row(table, 'Status of Support:', blank_other_check(row[36], "NASA Category"))
        add_table_row(table, 'Proposal/Award Number:', row[16])
        add_table_row(table, 'Source of Support:', blank_other_check(row[14], "Sponsor Name"))
        add_table_row(table, 'Primary Place of Performance:', blank_other_check(row[19], "Location"))
        add_table_row(table, 'Performance Period:', blank_other_check(row[17], 'Project Period'))
        add_table_row(table, 'Funding', currency_formatting(row[18], row[1]))
        add_table_row(table, 'Person Months:', blank_other_check(row[7],"Person Months"))       
        add_table_row(table, 'Objectives:', blank_other_check(row[4], "Goals"))
        add_table_row(table, 'Overlap:', blank_other_check(row[6], "Overlap"))     
               

    # Add an empty row for spacing
    table.add_row()

def fill_projects(sheet, starting_row, category_column, status, get_values_only=True):
    projects = []
    for row in sheet.iter_rows(min_row=starting_row, values_only=get_values_only):
        if row[category_column] == status:
            projects.append(row)
    return projects


def apply_paragraph_formatting(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Main code
version = "SPS-Common-Form 20251105"
print(20 * "*")
print(f"Version: {version}")
print("When reporting issues, please provide this version number")
#print("\nBe sure to enter the ID Number within this file\n")

print("\nNote: this program uses the NASA related columns in the spreadsheet.")

print(20 * "*")

# Select the 'C&P' sheet
workbook, workbook_file_path = open_file()
sheet = workbook['C&P']

# Get headers from row 40 on the spreadsheet
headers = create_column_dict(sheet, 40)

try:
    category_column = headers['NASA Category']  # set the department's category
except KeyError:
    print("Your category header has been changed -check to make sure it is properly spelled with no extra spaces")
    os.system(command="pause")
    exit()

# Create lists to store current and pending projects starting with row 41
#completed_projects = fill_projects(sheet, 41, category_column, 'Completed')
current_projects = fill_projects(sheet, 41, category_column, 'Current')
pending_projects = fill_projects(sheet, 41, category_column, 'Pending')




# Open the copied document
doc = docx.Document()



# Iterate through paragraphs and look for !-Start and !-End and remove any paragraphs between them.


#TODO remove repeating code

    # Add PI Name
pi_name = sheet.cell(4,2).value
try:
    if pi_name[:4] != "Name":
        print("There appears to be an issue with the top of the worksheet -No Name Found")
        print("Review the worksheet and ensure the instructions at the top are intact.")
        print("No File Saved")
        os.system(command="pause")
        exit()
except TypeError:
    print("There appears to be an issue with the top of the worksheet -No Name Found")
    print("Review the worksheet and ensure the instructions at the top are intact.")
    print("No File Saved")
    os.system(command="pause")
    exit()

name_paragraph = doc.add_paragraph()
run = name_paragraph.add_run("*")
run.bold = True
run.font.color.rgb = RGBColor(255,0,0)
run = name_paragraph.add_run("Name: ")
run.bold = True
run = name_paragraph.add_run(f'{pi_name[19:]}\n')
apply_paragraph_formatting(name_paragraph)

#Extract PID from A36
a36_value = sheet['A36'].value
# Extract text after the first colon, stripping any leading/trailing spaces
if a36_value and ":" in a36_value:
    pi_pid = a36_value.split(":", 1)[1].strip()
    if pi_pid:
        print(f"The PI's PID is {pi_pid}")
    else:
        print("No PID Entered -Please enter the PID manually")
else:
    print("Colon not found before PID")


# Add PID 
pid_paragraph = doc.add_paragraph()
run = pid_paragraph.add_run("Persistent Identifier (PID of the Senior/Key Person): ")
pid_paragraph.runs[0].bold = True
run = pid_paragraph.add_run(f"{pi_pid}\n")
apply_paragraph_formatting(pid_paragraph)

#Extract Position Title from B5
cell_b5_value = sheet['B5'].value
if cell_b5_value and "Title: " in cell_b5_value and " Start Date:" in cell_b5_value:
    start_index = cell_b5_value.find("Title: ") + len("Title: ")
    end_index = cell_b5_value.find(" Start Date:")
    pi_title = cell_b5_value[start_index:end_index].strip()
else:
    pi_title = ""  # or handle the case as needed
    print("There appears to be a problem with the position title, please add it manually to the document.")


# Add Position Title
position_paragraph = doc.add_paragraph()
run = position_paragraph.add_run("*")
run.bold = True 
run.font.color.rgb = RGBColor(255,0,0)
run = position_paragraph.add_run("Position Title: ")
run.bold = True
run = position_paragraph.add_run(f"{pi_title}\n")
apply_paragraph_formatting(position_paragraph)

# Add Organization and Location
org_paragraph = doc.add_paragraph("Organization and Location\n")
org_paragraph.runs[0].bold = True
org_paragraph.runs[0].underline = True
apply_paragraph_formatting(org_paragraph)

univ_paragraph = doc.add_paragraph()
run = univ_paragraph.add_run("*")
run.font.color.rgb = RGBColor(255,0,0)
run.bold = True
run = univ_paragraph.add_run("Name: Purdue University\n")
run.bold = True
apply_paragraph_formatting(univ_paragraph)

loc_paragraph = doc.add_paragraph()
run = loc_paragraph.add_run("*")
run.font.color.rgb = RGBColor(255,0,0)
run.bold = True
run = loc_paragraph.add_run("Location: West Lafayette, Indiana, United States\n")
run.bold = True
apply_paragraph_formatting(loc_paragraph)

# Add Proposals and Active Proposals
prop_paragraph = doc.add_paragraph("a.  Proposals and Active Proposals\n")
prop_paragraph.runs[0].bold = True
prop_paragraph.runs[0].underline = True
apply_paragraph_formatting(prop_paragraph)

# Add disclosure paragraph
disc_paragraph = doc.add_paragraph("In this section, disclose ALL proposals and active projects in accordance with the definition for ")
add_hyperlink(disc_paragraph, "http://nsf.gov/bfa/dias/policy/researchprotection/nspm33definitions.pdf", "current and pending (other) support")
apply_paragraph_formatting(disc_paragraph)

if current_projects:
    create_table(doc, current_projects, '')
    doc.add_paragraph()

if pending_projects:
    create_table(doc, pending_projects, '')


    
  


# Save the modified document
save_file(workbook_file_path)