#pyinstaller C:\SPS_python\.venv\Marshall\_dod_decimal.py --onefile
import os
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
import tkinter as tk
from tkinter import filedialog

filename = ''
v_effort_format_check = ''

# Open File Prompt
def open_file():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(
        filetypes=[("Excel files", "*.xlsx")],
        defaultextension=".xlsx",
        #initialdir=r"K:\_DeptAll\PreAward\3. Administrative\Faculty Documents - CPs\Faculty Current & Pending",
        title="SPS Current and Pending"
    )
    if not file_path:
        print("No File Selected")
        os.system="pause"
        exit()
    else:
        workbook_file_path = os.path.basename(file_path)
        return openpyxl.load_workbook(file_path), workbook_file_path
    
    #root.withdraw()  # Hide the main window

# Save File Prompt
def save_file(workbook_file_path):
    workbook_file_path = os.path.splitext(workbook_file_path)[0]
    save_path = filedialog.asksaveasfilename(
        filetypes=[("Word files", "*.docx")],
        title="SPS Current and Pending",
        #initialdir=r"K:\_DeptAll\PreAward\3. Administrative\Faculty Documents - CPs\Faculty Current & Pending",
        initialfile=f"{workbook_file_path}.docx"
    )
    
    if not save_path:
        print("No File Saved")
        os.system('pause')
        exit()

    # Check if the file has a .docx extension, if not, add it
    elif not save_path.lower().endswith('.docx'):
        save_path += '.docx'
        save_file_error_handling(save_path)
    else:
        save_file_error_handling(save_path)
        

def save_file_error_handling(save_path):
    try:
        document.save(save_path)
        print(f"Data has been exported to {save_path}")
        os.system('pause')

    except PermissionError:
        print("ERROR: File is open, close the file and try again")
        os.system('pause')

# Create a dictionary of column names and their corresponding column index
def create_column_dict(worksheet, row):
    column_dict = {}
    index = 0
    for cell in worksheet[row]:
        column_dict[cell.value] = index
        index+=1
    return column_dict

def effort_format_check():
    print("DARPA typically needs effort shown as Percent but some PIs want percent for non-DARPA Current & Pending")
    while True:
        sponsor_type= input("Did you need effort displayed as Decimal (D) or Percent (P)? ").lower()
        if sponsor_type in ['p', 'd']:
            return sponsor_type                          
        else:
            print("*" * 20)
            print("INVALID ENTRY")
       

# Function to create a table
def create_table(document, rows, table_title, effort_type):

    # Function to set font for a paragraph
    def set_font(paragraph, is_left_column=False):
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            if is_left_column:
                run.bold = True

    #Adds row to table and formats the text as needed
    def add_table_row(table, label, value, is_bold=False):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value) if value is not None else ''
        set_font(cells[0].paragraphs[0], is_left_column=True)
        set_font(cells[1].paragraphs[0])
        if is_bold:
            cells[1].paragraphs[0].runs[0].bold = True 
    
    #formats funding amount with dollar sign and two decimals. Handles error if the funding column has text
    def currency_formatting(funding_column,row):
        try:
            funding_float = float(funding_column if funding_column is not None else '')
            return f'${funding_float:,.2f}'
        except ValueError:
            print(f"The Funding Number is NOT text and says {funding_column} for title: {row}.\nEdit the Word file after saving or fix the Excel file and run this script again.")
            return '**** Incorrect Entry- Must be in format of ####.## ****'

    #Error checking to ensure other columns don't have blank spaces in needed rows.            
    def blank_other_check(column, label):
        if column is None:
            print(f"Something is blank that shouldn't be, check {label} column")
            print("No File Saved")
            os.system(command="pause")
            exit()
        elif column is not None and label == "Overlap":
            overlap_column = f"{column}\n\n"
            return overlap_column
        else:
            return column

    #Ensures the document header is the correct color, size, etc., regardless of style settings.
    title_formatting = document.add_paragraph()
    run = title_formatting.add_run(table_title)
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.size = Pt(14)
    run.bold = True
    
    #creates top two rows
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    table.style = 'Plain Table 4'
    #Set the width of the left column to 25%
    table.columns[0].width = Inches(1.5)
    table.columns[0].cells[0]._element.tcPr.tcW.type = 'pct'
    table.columns[0].cells[0]._element.tcPr.tcW.w = 1250  # 25% of 5000
    table.columns[1].width = Inches(8)

    # Set the preferred width of the right column to 75%
    table.columns[1].width = Inches(4.5)
    table.columns[1].cells[0]._element.tcPr.tcW.type = 'pct'
    table.columns[1].cells[0]._element.tcPr.tcW.w = 3750  # 75% of 5000

    #more text formatting and ensuring lines are single spaced
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = document.styles['Normal']
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    #Add a block of rows for each completed, current or pending proposal
    for row in rows:
        add_table_row(table, 'Title:', blank_other_check(row[1], "Title"))
        add_table_row(table, 'PI:', blank_other_check(row[13], "PI Name"))
        #add_table_row(table, 'Time Commitments:', blank_other_check(row[29],"DARPA Time Commitment"))
        if effort_type == 'p':
            add_table_row(table, 'Time Commitments:', blank_other_check(row[29],"DARPA Time Commitment"))
        else:
            add_table_row(table, 'Time Commitments:', blank_other_check(row[30],"DOD Time Commitment"))
        add_table_row(table, 'Agency:', blank_other_check(row[14], "Sponsor Name"))
        add_table_row(table, 'Agency Address:', row[20])
        add_table_row(table, "Agency's Contact/Contracting Grants Office:", row[5])
        add_table_row(table, 'Performance Period:', blank_other_check(row[17], 'Project Period'))
        add_table_row(table, 'Funding', currency_formatting(row[18], row[1]))
        add_table_row(table, 'Objectives:', blank_other_check(row[4], "Goals"))
        add_table_row(table, 'Overlap:', blank_other_check(row[6], "Overlap"))

    # Add an empty row for spacing
    table.add_row()

#pulls data from spreadsheet and puts it into a list for use by create_table
def fill_projects(sheet, starting_row, category_column, status, get_values_only=True):
    projects = []

    for row in sheet.iter_rows(min_row=starting_row, values_only=get_values_only):
        if row[category_column] == status:
            projects.append(row)
    
    return projects

# creates paragraph above table
def create_document():
    # Create a new Word document
    document = Document()
    styles = document.styles
    style = styles.add_style('Plain Table 4', WD_STYLE_TYPE.TABLE)
    # style = document.styles['Normal'] to remove
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add a title to the document
    pi_name = sheet.cell(4,2).value
    try:
        if pi_name[:4] != "Name":
            print("There appears to be an issue with the top of the worksheet -No Name Found")
            print("Review the worksheet and ensure the instructions at the top are intact.")
            print("No File Saved")
            os.system(command="pause")
            exit()
    except TypeError:
        print("There appears to be an issue with the top of the worksheet -No Name Found")
        print("Review the worksheet and ensure the instructions at the top are intact.")
        print("No File Saved")
        os.system(command="pause")
        exit()


    document.add_paragraph(pi_name)
    document.add_paragraph("Commons ID:")
    other_support_text = document.add_paragraph()
    run = other_support_text.add_run("Other Support - Project/Proposal")
    run.bold = True
    run.font.size = Pt(14)

    return document

#creates paragraphs below tables
def create_in_kind_page(document):
    #add page break before In-Kind
    need_inkind = input("Did you need an In-Kind section? 'Y' or 'N' ").lower()
    if need_inkind == 'y':
        document.add_page_break()

        in_kind_top = document.add_paragraph()
        in_kind_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = in_kind_top.add_run("\n\nIn-Kind")
        run.font.size = Pt(14)
        run.bold = True

        document.add_paragraph(
        f'\n\nSummary of In-Kind Contribution: None\n*Status of Support: \n *Source of Support: \n Project/Proposal Start & End Date: (MM/YY)\n*Person Months (Calendar, Academic, Summer) per budget period:\n\n*Estimated Dollar Value of In-Kind Information:\n\n *Overlap (summarized for each individual)\n\n')
        document.add_paragraph(f'I, PD/PI or other senior/key personnel, certify that the statements herein are true, complete, and accurate to the best of my knowledge, and accept the obligation to complete with Public Health Services terms and conditions if a grant is awarded as a result of this application. I am aware that any false, fictitious or fradulent statements or claims may subject me to criminal, civil, or administrative penalties.')
        document.add_paragraph("*Signature:")
        document.add_paragraph("Date: ")
    else:
        print("You did not select Y so no In-Kind section was added.")

# Main code
version = "SPS-DOD-20250325"
print(20 * "*")
print(f"Version: {version}")
print("When reporting issues, please provide this version number")
print(20 * "*")

v_effort_format_check = effort_format_check()
print(f"It is {v_effort_format_check}")

# Load the Excel workbook
workbook, workbook_file_path = open_file()

# Select the 'C&P' sheet
sheet = workbook['C&P']

# Get headers from row 40 on the spreadsheet
headers = create_column_dict(sheet, 40)

try:
    category_column = headers['AFOSR/DARPA/DOD Category'] #set the department's category
except KeyError:
    print("Your category header has been changed -check to make sure it is properly spelled with no extra spaces")
    os.system(command="pause")
    exit()

# Create lists to store current and pending projects starting with row 41
completed_projects = fill_projects(sheet, 41, category_column, 'Completed')
current_projects = fill_projects(sheet, 41, category_column, 'Current')
pending_projects = fill_projects(sheet, 41, category_column, 'Pending')

# Create a new Word document
document = create_document()

# Create a table for projects
if completed_projects == []:
    print("No completed projects to show. Skipped")
else:
    create_table(document, completed_projects, 'Completed Projects', v_effort_format_check )
    paragraph = document.add_paragraph()

if current_projects == []:
    print("No current projects to show. Skipped")
else:
    create_table(document, current_projects, 'Current Projects', v_effort_format_check )
    paragraph = document.add_paragraph()

if pending_projects == []:
    print("No pending projects to show. Skipped")
else:
    create_table(document, pending_projects, 'Pending Projects', v_effort_format_check)
    paragraph = document.add_paragraph()

# Create in kind page
create_in_kind_page(document)

# Save the file
save_file(workbook_file_path)