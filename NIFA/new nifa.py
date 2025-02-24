import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_UNDERLINE

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog
import os

# Open File Prompt
def open_file():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    file_path = filedialog.askopenfilename(
        filetypes=[("Excel files", "*.xlsx")],
        defaultextension=".docx"
    )

    if not file_path:
        print("No File Selected")
    else:
        return openpyxl.load_workbook(file_path)

# Save File Prompt
def save_file():
    save_path = filedialog.asksaveasfilename(
        filetypes=[("Word files", "*.docx")]
    )
    
    if not save_path:
        print("No File Saved")
        os.system('pause')

    # Check if the file has a .docx extension, if not, add it
    elif not save_path.lower().endswith('.docx'):
        save_path += '.docx'
        save_file_error_handling(save_path)
    else:
        save_file_error_handling(save_path)
        

def save_file_error_handling(save_path):
    try:
        doc.save(save_path)
        print(f"Data has been exported to {save_path}")
        os.system('pause')

    except PermissionError:
        print("ERROR: File is open, close the file and try again")
        os.system('pause')

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 12, "val": "single", "color": "#000000", "space": "0"}
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for side, attrs in kwargs.items():
        pBdr = OxmlElement('w:tcBorders')
        tcPr.append(pBdr)
        
        border = OxmlElement(f'w:{side}')
        pBdr.append(border)

        for key, value in attrs.items():
            border.set(qn(f'w:{key}'), str(value))



# Load the Excel file
#workbook = openpyxl.load_workbook(r'.venv\NIH C_P Olson Matther - Final.xlsx')
workbook = open_file()
sheet = workbook['C&P']

# Get the headers from row 40
headers = [cell.value for cell in sheet[40]]

# Create lists for active and pending projects
active_projects = []
pending_projects = []

# Iterate through rows starting from row 41
for row in sheet.iter_rows(min_row=41, values_only=True):
    nifa_category = row[headers.index('NIFA/AFRI/USDA Category')]
    if nifa_category == 'Active':
        active_projects.append((
            row[headers.index('NIFA/AFRI/USDA Person Name')],
            row[headers.index('NIFA/AFRI/USDA Agency Source')],
            row[headers.index('Total Award Amount (including Indirect Costs):  ')],
            row[headers.index('Project Period')],
            row[headers.index('Percentage of Time Committed')],
            row[headers.index('Project Title')]
        ))
    elif nifa_category == 'Pending':
        pending_projects.append((
            row[headers.index('NIFA/AFRI/USDA Person Name')],
            row[headers.index('NIFA/AFRI/USDA Agency Source')],
            row[headers.index('Total Award Amount (including Indirect Costs):  ')],
            row[headers.index('Project Period')],
            row[headers.index('Percentage of Time Committed')],
            row[headers.index('Project Title')]
        ))

# Function to create table in Word document
def create_table(doc, projects, title, category):

    #formats funding amount with dollar sign and two decimals. Handles error if the funding column has text
    def currency_formatting(funding_column):
        try:
            funding_float = float(funding_column if funding_column is not None else '')
            return f'${funding_float:,.2f}'
        except ValueError:
            print(f"The Funding Number is NOT text and says {funding_column}.\n Edit the Word file after saving or fix the Excel file and run this script again.")
            return '**** Incorrect Entry- Must be in format of ####.## ****'
            

    table = doc.add_table(rows=2, cols=6)
    table.autofit = False
    table.style = 'Plain Table 4'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set custom borders for all cells
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # Remove left border for all cells
            set_cell_border(
                cell,
                left={"sz": 0, "val": "none", "color": "#000000", "space": "0"}
            )
            
            # Add right border for all cells except the last column
            if j < len(row.cells) - 1:
                set_cell_border(
                    cell,
                    right={"sz": 4, "val": "single", "color": "#000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell,
                    right={"sz": 0, "val": "none", "color": "#000000", "space": "0"}
                )
            
            # Keep the existing top and bottom border logic
            if i < 2:
                set_cell_border(
                    cell,
                    top={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                    #bottom={"sz": 4, "val": "single", "color": "#000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell,
                    top={"sz": 0, "val": "none", "color": "#000000", "space": "0"},
                    bottom={"sz": 0, "val": "none", "color": "#000000", "space": "0"}
                )
    
    # Set header row
    hdr_cells = table.rows[0].cells
    headers = ['Person Name', 'Agency Source', 'Total Amount', 'Dates', ' % Time Committed', 'Title']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set category row
    category_cells = table.rows[1].cells
    category_cells[0].text = f"{category}\n"
    for i in range(1, 6):
        category_cells[i].text = ""
    
    # Make the category row bold
    for cell in table.rows[1].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    
    # Add project data
    for index, (person, agency, amount, dates, time, title) in enumerate(projects):
        row_cells = table.add_row().cells
        for j, cell in enumerate(row_cells):
            # Set custom borders for new rows
            set_cell_border(
                cell,
                left={"sz": 0, "val": "none", "color": "#000000", "space": "0"},
                top={"sz": 0, "val": "none", "color": "#000000", "space": "0"}
            )
            # Add right border for all cells except the last column
            if j < len(row_cells) - 1:
                set_cell_border(
                    cell,
                    right={"sz": 4, "val": "single", "color": "#000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell,
                    right={"sz": 0, "val": "none", "color": "#000000", "space": "0"}
                )
            
            # Add bottom border only for the last row
            if index == len(projects) - 1:
                set_cell_border(
                    cell,
                    bottom={"sz": 4, "val": "single", "color": "#000000", "space": "0"}
                )
            else:
                set_cell_border(
                    cell,
                    bottom={"sz": 0, "val": "none", "color": "#000000", "space": "0"}
                )

        row_cells[0].text = str(person)
        row_cells[1].text = str(agency)
        row_cells[2].text = str(currency_formatting(amount))
        row_cells[3].text = str(dates)
        row_cells[4].text = str(time)
        row_cells[5].text = str(title)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def create_document():
    # Create a new Word document
    document = Document()
    styles = document.styles
    style = styles.add_style('Plain Table 4', WD_STYLE_TYPE.TABLE)
    # style = document.styles['Normal'] to remove
    style.font.name = 'Times New Roman'
    style.font.size = Pt(7.5)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add a title to the document
  
    # Add a table with 1 row and 1 column
    header_table = document.add_table(rows=1, cols=1)    

    # Get the cell and add the text
    cell = header_table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("\nCurrent & Pending Support")

    set_cell_border(
                    cell,
                    top={"sz": 18, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 18, "val": "single", "color": "#000000", "space": "0"},
                    left={"sz": 18, "val": "single", "color": "#000000", "space": "0"},
                    right={"sz": 18, "val": "single", "color": "#000000", "space": "0"}
                )

    # Center the text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set the font to Times New Roman, 14pt
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True

    document.add_paragraph()

    info_table = document.add_table(rows=2, cols=1)

    # Fill the first row with "Name"
    pi_text = sheet.cell(4,2).value
    pi_name = pi_text.split(": ")[1]
    cell= info_table.cell(0, 0)
    cell.text = f"Name: {pi_name}"
    run = cell.paragraphs[0].runs[0]
    run.bold=True


    # Fill the second row with "Instructions"
    sec_cell= info_table.cell(1, 0)
 
    sec_cell.text = "Instruct"

    # Get the second row of the info_table so you can set the border correctly
    second_row = info_table.rows[1]

    # Add top border to each cell in the second row
    for cell in second_row.cells:
        set_cell_border(
            cell,
            top={"sz": 4, "val": "single", "color": "#000000", "space": "0"}
        )

    return document

# Create Word document and add tables
doc = create_document()
create_table(doc, active_projects, 'Active Projects', 'Active')
#doc.add_page_break()
create_table(doc, pending_projects, 'Pending Projects', 'Pending')

# Save the Word document
save_file()
