#pyinstaller .venv\nih.py --onefile
import os
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
import tkinter as tk
from tkinter import filedialog


# Open File Prompt
def open_file():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    file_path = filedialog.askopenfilename(
        filetypes=[("Excel files", "*.xlsx")],
        defaultextension=".docx"
    )

    if not file_path:
        print("No File Selected")
    else:
        return openpyxl.load_workbook(file_path)

# Save File Prompt
def save_file():
    save_path = filedialog.asksaveasfilename(
        filetypes=[("Word files", "*.docx")]
    )
    
    if not save_path:
        print("No File Saved")
        os.system('pause')

    # Check if the file has a .docx extension, if not, add it
    elif not save_path.lower().endswith('.docx'):
        save_path += '.docx'
        save_file_error_handling(save_path)
    else:
        save_file_error_handling(save_path)
        

def save_file_error_handling(save_path):
    try:
        document.save(save_path)
        print(f"Data has been exported to {save_path}")
        os.system('pause')

    except PermissionError:
        print("ERROR: File is open, close the file and try again")
        os.system('pause')



# Create a dictionary of column names and their corresponding column index
def create_column_dict(worksheet, row):
    column_dict = {}
    index = 0
    for cell in worksheet[row]:
        column_dict[cell.value] = index
        index+=1
    return column_dict

# Function to create a table
def create_table(document, rows, table_title):
    # Function to set font for a paragraph
    def set_font(paragraph, is_left_column=False):
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            if is_left_column:
                run.bold = True

    def add_table_row(table, label, value, is_bold=False):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value) if value is not None else ''
        set_font(cells[0].paragraphs[0], is_left_column=True)
        set_font(cells[1].paragraphs[0])
        if is_bold:
            cells[1].paragraphs[0].runs[0].bold = True 
    
    #formats funding amount with dollar sign and two decimals. Handles error if the funding column has text
    def currency_formatting(funding_column):
        try:
            funding_float = float(funding_column if funding_column is not None else '')
            return f'${funding_float:,.2f}'
        except ValueError:
            print(f"The Funding Number is NOT text and says {funding_column}.\n Edit the Word file after saving or fix the Excel file and run this script again.")
            return '**** Incorrect Entry- Must be in format of ####.## ****'

    #concacts mini headings over yearly effort    
    def person_month_formatting(effort_column):
       return "Year  Person Months (##.##)\n"  + str(effort_column if effort_column is not None else '')

    title_formatting = document.add_paragraph()
    run = title_formatting.add_run(table_title)
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.size = Pt(14)
    run.bold = True
    #document.add_heading(table_title, level=1) to remove
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    table.style = 'Plain Table 4'
    #Set the width of the left column to 25%
    table.columns[0].width = Inches(1.5)
    table.columns[0].cells[0]._element.tcPr.tcW.type = 'pct'
    table.columns[0].cells[0]._element.tcPr.tcW.w = 1250  # 25% of 5000
    table.columns[1].width = Inches(8)

    # Set the preferred width of the right column to 75%
    table.columns[1].width = Inches(4.5)
    table.columns[1].cells[0]._element.tcPr.tcW.type = 'pct'
    table.columns[1].cells[0]._element.tcPr.tcW.w = 3750  # 75% of 5000

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = document.styles['Normal']
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for row in rows:
        add_table_row(table, 'Title:', row[1])
        add_table_row(table, 'Major Goals:', row[4])
        add_table_row(table, 'Status of Support:', row[25], is_bold=True)
        add_table_row(table, 'Project Number:', row[26])
        add_table_row(table, 'Name of PD/PI:', row[13])
        add_table_row(table, 'Prime Sponsor:', row[15])
        add_table_row(table, 'Source of Support:', row[14])
        add_table_row(table, 'Primary Place of Performance:', row[19])
        add_table_row(table, 'Project/Proposal Start & End Date:', row[17])
        add_table_row(table, 'Funding', currency_formatting(row[18]))
        add_table_row(table, '*Person Months:', person_month_formatting(row[7]))

    # Add an empty row for spacing
    table.add_row()

def fill_projects(sheet, starting_row, category_column, status, get_values_only=True):
    projects = []

    for row in sheet.iter_rows(min_row=starting_row, values_only=get_values_only):
        if row[category_column] == status:
            projects.append(row)
    
    return projects

def create_document():
    # Create a new Word document
    document = Document()
    styles = document.styles
    style = styles.add_style('Plain Table 4', WD_STYLE_TYPE.TABLE)
    # style = document.styles['Normal'] to remove
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add a title to the document
    #document.add_heading('NIH Current and Pending Report', 0) to remove
    pi_name = sheet.cell(4,2).value
    try:
        if pi_name[:4] != "Name":
            print("There appears to be an issue with the top of the worksheet -No Name Found")
            print("Review the worksheet and ensure the instructions at the top are intact.")
            print("No File Saved")
            os.system(command="pause")
            exit()
    except TypeError:
        print("There appears to be an issue with the top of the worksheet -No Name Found")
        print("Review the worksheet and ensure the instructions at the top are intact.")
        print("No File Saved")
        os.system(command="pause")
        exit()


    document.add_paragraph(pi_name)
    document.add_paragraph("Commons ID:")
    other_support_text = document.add_paragraph()
    run = other_support_text.add_run("Other Support - Project/Proposal")
    run.bold = True
    run.font.size = Pt(14)

    return document

def create_in_kind_page(document):
    #add page break before In-Kind
    document.add_page_break()

    in_kind_top = document.add_paragraph()
    in_kind_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = in_kind_top.add_run("\n\nIn-Kind")
    run.font.size = Pt(14)
    run.bold = True

    document.add_paragraph(
    f'\n\nSummary of In-Kind Contribution: None\n*Status of Support: \n *Source of Support: \n Project/Proposal Start & End Date: (MM/YY)\n*Person Months (Calendar, Academic, Summer) per budget period:\n\n*Estimated Dollar Value of In-Kind Information:\n\n *Overlap (summarized for each individual)\n\n')
    document.add_paragraph(f'I, PD/PI or other senior/key personnel, certify that the statements herein are true, complete, and accurate to the best of my knowledge, and accept the obligation to complete with Public Health Services terms and conditions if a grant is awarded as a result of this application. I am aware that any false, fictitious or fradulent statements or claims may subject me to criminal, civil, or administrative penalties.')
    document.add_paragraph("*Signature:")
    document.add_paragraph("Date: ")

# Main code
# Load the Excel workbook
workbook = open_file()

# Select the 'C&P' sheet
sheet = workbook['C&P']

# Get headers from row 40 on the spreadsheet
headers = create_column_dict(sheet, 40)

try:
    category_column = headers['NIH Category'] #set the department's category
except KeyError:
    print("Your category header has been changed -check to make sure it is properly spelled with no extra spaces")
    os.system(command="pause")
    exit()
# Create lists to store current and pending projects starting with row 41
current_projects = fill_projects(sheet, 41, category_column, 'Awarded')
pending_projects = fill_projects(sheet, 41, category_column, 'Pending')

# Create a new Word document
document = create_document()

# Create a table for projects
create_table(document, current_projects, 'Awarded Projects')
create_table(document, pending_projects, 'Pending Projects')
# Create a table for projects

if current_projects == []:
    print("No current projects to show. Skipped")
else:
    create_table(document, current_projects, 'Current Projects')

if pending_projects == []:
    print("No pending projects to show. Skipped")
else:
    create_table(document, pending_projects, 'Pending Projects')

# Create in kind page
create_in_kind_page(document)

# Save the file
save_file()